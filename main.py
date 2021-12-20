from datetime import date
import docx
import os
from docx.shared import Pt
from docx.enum.style import WD_STYLE_TYPE

doc = docx.Document()

today = date.today()



DATE = str(today)
COMPANY = input("\nEnter the company: ")
POSITION = input("Enter the position: ")

try:
    os.makedirs("cover_letters")
except:
    pass


cover_letter = ""


parag = doc.add_paragraph()

obj_styles = doc.styles
obj_charstyle = obj_styles.add_style('CommentsStyle', WD_STYLE_TYPE.CHARACTER)
obj_font = obj_charstyle.font
obj_font.size = Pt(15)
obj_font.name = 'Times New Roman'

print(cover_letter)


print("\nSuccessful!")

parag.add_run(cover_letter, style='CommentsStyle').bold = True

doc.save("cover_letters/" + COMPANY + ".docx")